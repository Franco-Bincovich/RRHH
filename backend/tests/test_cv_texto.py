"""
`_cv_texto.extraer`: del archivo del CV a texto plano.

## 🔴 LOS PDF SON PDF DE VERDAD, GENERADOS EN EL TEST

No hay un solo mock de `pypdf` acá. Un doble que devolviera `"texto de prueba"` no puede
desmentir nada de lo que importa: si el parser está bien invocado, si un archivo cifrado se
detecta, si uno truncado revienta o se reporta. Los archivos se construyen con `reportlab` (que
ya estaba, para el export de PDF) y con `pypdf` para el cifrado, así que lo que se prueba es el
comportamiento real de la librería sobre bytes reales.

## 🔴 EL PRINCIPIO QUE ORDENA TODO EL MÓDULO: NUNCA LEVANTA

Cifrado, corrupto, sin texto, formato no soportado — los cuatro devuelven un `warning`, no una
excepción. Si alguno propagara, se llevaría puesta la creación del candidato, y la postulación se
perdería por un problema del ARCHIVO y no de la persona.

`TestElLoteSigue` lo verifica donde importa de verdad: en el alta, con DOS CVs en el mismo mail,
uno bueno y uno roto. **Con uno solo, "el lote sigue" y "el lote se cortó" son indistinguibles.**
"""
import os

_TEST_ENV: dict[str, str] = {
    "SUPABASE_URL": "https://test-project.supabase.co",
    "SUPABASE_ANON_KEY": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.test.anon",
    "SUPABASE_SERVICE_KEY": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.test.service",
    "JWT_SECRET": "test-secret-for-unit-tests-only-minimum-32-chars!!",
    "ANTHROPIC_API_KEY": "sk-ant-test",
    "RESEND_API_KEY": "re_test",
}
for _k, _v in _TEST_ENV.items():
    os.environ.setdefault(_k, _v)

import io  # noqa: E402
from types import SimpleNamespace  # noqa: E402
from uuid import uuid4  # noqa: E402

import pytest  # noqa: E402

from services._cv_alta import crear_de_un_cv  # noqa: E402
from services._cv_texto import MAX_CARACTERES, MINIMO_UTIL, extraer  # noqa: E402

_PARRAFO = ("Analista de datos con ocho anios de experiencia en Python y SQL. "
            "Licenciado en Sistemas por la UBA. Ingles avanzado. ")


# ── constructores de archivos REALES ──────────────────────────────────────────

def _pdf(texto: str, paginas: int = 1) -> bytes:
    """Un PDF real con texto, hecho con reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for _ in range(paginas):
        y = 800
        # Líneas cortas: reportlab no hace wrap solo y el texto de más se perdería.
        for i in range(0, len(texto), 90):
            c.drawString(40, y, texto[i:i + 90])
            y -= 12
            if y < 40:
                break
        c.showPage()
    c.save()
    return buf.getvalue()


def _pdf_sin_texto() -> bytes:
    """Un PDF con UNA FIGURA y ningún texto: es lo que produce un escaneo."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.rect(50, 50, 400, 600, fill=1)
    c.showPage()
    c.save()
    return buf.getvalue()


def _pdf_cifrado(password: str = "secreto") -> bytes:
    """Un PDF real cifrado con contraseña de usuario."""
    from pypdf import PdfReader, PdfWriter

    w = PdfWriter()
    for p in PdfReader(io.BytesIO(_pdf(_PARRAFO * 4))).pages:
        w.add_page(p)
    w.encrypt(password)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def _docx(parrafos: list, filas_tabla: list = None) -> bytes:
    from docx import Document

    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    if filas_tabla:
        t = doc.add_table(rows=len(filas_tabla), cols=len(filas_tabla[0]))
        for i, fila in enumerate(filas_tabla):
            for j, celda in enumerate(fila):
                t.rows[i].cells[j].text = celda
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 1. PDF ────────────────────────────────────────────────────────────────────

class TestPdf:

    def test_extrae_el_texto_de_un_pdf_real(self) -> None:
        """¿Qué tendría que ser distinto para que falle? Que el PDF fuera un mock: ahí el test
        estaría comprobando lo que el propio fake devuelve, no que pypdf esté bien invocado."""
        r = extraer(_pdf(_PARRAFO * 4), "cv.pdf")
        assert r.warning is None
        assert "Analista de datos" in r.texto and "UBA" in r.texto

    def test_un_pdf_sin_texto_avisa_y_no_revienta(self) -> None:
        """El escaneo: hay páginas, hay tinta, no hay capa de texto."""
        r = extraer(_pdf_sin_texto(), "cv.pdf")
        assert r.texto is None and "texto extraíble" in r.warning

    def test_un_pdf_cifrado_avisa_la_contrasena(self) -> None:
        """🔴 El motivo tiene que ser ESE y no "corrupto": la acción de RRHH es distinta —pedirle
        la contraseña, no el archivo de nuevo—. Es la razón por la que el warning es texto."""
        r = extraer(_pdf_cifrado(), "cv.pdf")
        assert r.texto is None and "contraseña" in r.warning

    def test_un_pdf_corrupto_avisa_y_no_revienta(self) -> None:
        """Truncado a la mitad: pypdf levanta `PdfStreamError` y tiene que quedar contenido."""
        entero = _pdf(_PARRAFO * 4)
        r = extraer(entero[: len(entero) // 2], "cv.pdf")
        assert r.texto is None and "corrupto" in r.warning

    def test_bytes_que_no_son_un_pdf(self) -> None:
        assert extraer(b"esto no es un pdf", "cv.pdf").texto is None

    def test_un_pdf_de_varias_paginas_junta_todas(self) -> None:
        r = extraer(_pdf(_PARRAFO, paginas=3), "cv.pdf")
        assert r.texto.count("Analista de datos") == 3


# ── 2. DOCX y .doc ────────────────────────────────────────────────────────────

class TestDocx:

    def test_extrae_parrafos(self) -> None:
        r = extraer(_docx([_PARRAFO * 2, "Experiencia laboral detallada." * 5]), "cv.docx")
        assert r.warning is None and "Analista de datos" in r.texto

    def test_tambien_lee_las_tablas(self) -> None:
        """🔴 En un CV media hoja de vida vive en una tabla. Sin esto, un CV maquetado con tablas
        —que es lo normal en las plantillas de Word— daría "sin texto extraíble" teniéndolo todo.

        ¿Qué tendría que ser distinto para que falle? Que el docx tuviera el texto también en
        párrafos: acá el contenido está SOLO en la tabla, así que si no se leyera no llegaría a
        `MINIMO_UTIL` y saldría warning."""
        r = extraer(_docx([], [["Puesto", "Empresa"], [_PARRAFO * 2, _PARRAFO]]), "cv.docx")
        assert r.warning is None and "Analista de datos" in r.texto

    def test_un_docx_corrupto_avisa(self) -> None:
        assert extraer(b"PK\x03\x04 roto", "cv.docx").warning is not None


def test_el_doc_viejo_dice_que_no_se_soporta() -> None:
    """🔴 El motivo nombra la acción: "pedile el CV en PDF o DOCX". Un "no se pudo procesar"
    genérico deja a RRHH abriendo el archivo para descubrir que era un .doc."""
    r = extraer(b"\xd0\xcf\x11\xe0cualquier cosa", "cv.doc")
    assert r.texto is None
    assert "no soportado" in r.warning and "PDF o DOCX" in r.warning


def test_una_extension_desconocida_no_revienta() -> None:
    assert extraer(b"...", "cv.rtf").warning is not None


# ── 3. Los topes ──────────────────────────────────────────────────────────────

class TestTopes:

    def test_trunca_al_maximo_y_lo_avisa(self) -> None:
        """🔴 Se trunca, NO se descarta: un CV recortado sigue siendo clasificable. Y se avisa,
        porque quien lee la ficha tiene que saber que el sistema no vio el final.

        ¿Qué tendría que ser distinto para que falle? Que el PDF fuera más corto que el tope —
        con un CV normal esta rama no se ejecuta y el truncado quedaría sin probar."""
        largo = _pdf(_PARRAFO * 40, paginas=6)
        r = extraer(largo, "cv.pdf")
        assert len(r.texto) == MAX_CARACTERES
        assert "muy largo" in r.warning

    def test_debajo_del_minimo_util_es_warning_no_texto(self) -> None:
        """Un PDF con tres palabras no alcanza para clasificar. El piso existe porque un escaneo
        devuelve basura corta (un número de página), no vacío — "hay algo" no sirve de criterio."""
        r = extraer(_pdf("Juan Perez"), "cv.pdf")
        assert r.texto is None and len("Juan Perez") < MINIMO_UTIL

    def test_justo_por_encima_del_minimo_si_pasa(self) -> None:
        """El control del test anterior: el piso corta abajo, no arriba."""
        r = extraer(_pdf("a" * 40 + " " + _PARRAFO * 3), "cv.pdf")
        assert r.texto is not None and len(r.texto) >= MINIMO_UTIL

    def test_el_texto_sale_normalizado(self) -> None:
        """Saltos de línea y espacios repetidos colapsan: son ruido que se paga en tokens."""
        r = extraer(_pdf(_PARRAFO * 3), "cv.pdf")
        assert "\n" not in r.texto and "  " not in r.texto


# ── 4. 🔴 Un archivo roto NO impide que el candidato se cree ──────────────────

class _CandidatoRepo:
    def __init__(self) -> None:
        self.creados: list = []

    def existe_cv_de_gmail(self, *a) -> bool:
        return False

    def save_candidato(self, vacante_id, data, empresa_id, origen=None):
        c = SimpleNamespace(id=f"cand-{len(self.creados) + 1}", **(origen or {}))
        self.creados.append(c)
        return c

    def set_cv(self, cid, path) -> None:
        pass


class _CvService:
    def subir(self, *a) -> str:
        return "path/cv"


class TestElLoteSigue:
    """🔴 DOS CVs en el mismo mail: uno bueno y uno roto.

    Con UNO SOLO, "el lote sigue" y "el lote se cortó" son indistinguibles — las dos
    implementaciones dejarían 0 o 1 candidato y ninguna aserción podría separarlas. Con dos, la
    correcta deja 2 y la rota deja 1 (o levanta).
    """

    def _cv(self, contenido: bytes, filename: str):
        return SimpleNamespace(contenido=contenido, filename=filename, mime="application/pdf",
                               attachment_id="a1")

    def _correr(self, repo, cvs) -> None:
        vacante = SimpleNamespace(id=str(uuid4()), empresa_id=str(uuid4()))
        res = SimpleNamespace(message_id="m1", candidatos_creados=[], ya_existian=0)
        for cv in cvs:
            crear_de_un_cv(res, cv, vacante, {"From": "Ana <a@x.com>"},
                           candidato_repo=repo, cv_service=_CvService())

    def test_el_cv_roto_no_se_lleva_al_bueno(self) -> None:
        repo = _CandidatoRepo()
        self._correr(repo, [self._cv(b"roto", "roto.pdf"),
                            self._cv(_pdf(_PARRAFO * 4), "bueno.pdf")])
        assert len(repo.creados) == 2, "un archivo ilegible cortó el lote"

    def test_el_roto_queda_con_warning_y_sin_texto(self) -> None:
        repo = _CandidatoRepo()
        self._correr(repo, [self._cv(b"roto", "roto.pdf"),
                            self._cv(_pdf(_PARRAFO * 4), "bueno.pdf")])
        roto, bueno = repo.creados
        assert roto.cv_texto is None and roto.screening_warning is not None
        assert bueno.cv_texto is not None and bueno.screening_warning is None

    def test_el_texto_viaja_en_el_MISMO_insert(self) -> None:
        """No hay un update posterior: si lo hubiera, un fallo entre el INSERT y el UPDATE dejaría
        el candidato sin texto y sin warning, o sea indistinguible de uno sin CV."""
        repo = _CandidatoRepo()
        self._correr(repo, [self._cv(_pdf(_PARRAFO * 4), "cv.pdf")])
        assert hasattr(repo.creados[0], "cv_texto")

    @pytest.mark.parametrize("filename", ["cv.doc", "cv.pdf"], ids=["doc", "pdf_roto"])
    def test_ningun_formato_problematico_levanta(self, filename) -> None:
        repo = _CandidatoRepo()
        self._correr(repo, [self._cv(b"contenido invalido", filename)])
        assert len(repo.creados) == 1


def test_el_warning_llega_al_schema() -> None:
    """🔴 El mapper podría traer la columna y el schema descartarla EN SILENCIO — pasó tres veces
    en este repo. Se verifica contra el mapper REAL, no contra un fake."""
    from datetime import datetime, timezone

    from repositories._candidato_row import _crow
    from schemas.candidato import CandidatoResponse

    fila = {"id": str(uuid4()), "vacante_id": None, "empresa_id": str(uuid4()), "nombre": "Ana",
            "apellido": "Pérez", "email": "a@x.com", "etapa": "postulado",
            "screening_warning": "Formato .doc no soportado.",
            "created_at": datetime.now(timezone.utc)}
    assert _crow(fila).screening_warning == "Formato .doc no soportado."
    assert "screening_warning" in CandidatoResponse.model_fields
    # `cv_texto` NO se expone: es la entrada del clasificador, no un dato de pantalla.
    assert "cv_texto" not in CandidatoResponse.model_fields
