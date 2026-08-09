"""
Del archivo del CV a texto plano. Fase 2 de 3 del CV screening: es la ENTRADA del clasificador.

Función pura sobre bytes: no toca la red, ni la base, ni Storage. Devuelve siempre un resultado
—nunca levanta— porque un archivo que no se puede leer **no es un error del lote**: es un
candidato al que hay que pedirle el CV de otra forma.

## 🔴 NUNCA LEVANTA, Y ESO NO ES TRAGARSE EL ERROR

Un PDF puede venir cifrado, corrupto, con 0 páginas, o ser un escaneo sin capa de texto. Si
cualquiera de esos tumbara la extracción, se llevaría puesta la creación del candidato — y la
postulación se perdería por un problema del ARCHIVO, no de la persona. El motivo no se pierde: va
a `screening_warning`, que es texto y le dice a RRHH qué pedirle al candidato. Mismo criterio que
ya rige para la validación del adjunto y para el fallo de Storage.

## 🔴 EL TOPE DE 20.000 CARACTERES, Y POR QUÉ HAY QUE PONERLE UNO

Esto va a viajar a Claude en la fase 3, y **los tokens se pagan por caracter**. Un CV normal son
2.000–6.000 caracteres; uno de 15 páginas con historia académica completa pasa los 40.000. Sin
tope, un solo archivo largo puede costar varias veces lo que costó todo el lote, y el excedente
casi nunca aporta: lo que decide una preselección está en las primeras páginas.

20.000 caracteres son ~5.000 tokens, holgado para un CV largo de verdad y muy por debajo de la
ventana del modelo. **Se trunca, no se descarta**, y el truncado se AVISA: un CV recortado sigue
siendo clasificable, pero quien lea la ficha tiene que saber que el sistema no vio el final.

⚠️ Es constante de módulo y NO variable de entorno, mismo criterio que `LIMITE_FILAS_EXPORT`:
subirlo es una decisión de costo, no configuración.

## Formatos

  · **PDF** → `pypdf`. Puro Python, sin binarios nativos (ver el porqué en `requirements.txt`).
  · **DOCX** → `python-docx`, que ya estaba para el export.
  · **DOC viejo** (OLE2, pre-2007) → NO se soporta. Leerlo pide `antiword` (binario) o `textract`
    (cadena de deps pesada) para un formato que Word exporta a PDF en dos clicks. Se marca y listo.

⚠️ **Sin OCR, a propósito.** Un PDF escaneado se marca "sin texto extraíble". Tesseract son ~50 MB
más binarios nativos, segundos por página, y resuelve un caso que RRHH resuelve abriendo el
archivo. La decisión se revisa cuando se sepa qué proporción de los CVs reales son escaneos.
"""
import io
from dataclasses import dataclass
from typing import Optional

from utils.logger import logger

# Ver el bloque del encabezado. ~5.000 tokens.
MAX_CARACTERES = 20_000

# Debajo de esto, lo extraído no alcanza para clasificar: un PDF escaneado suele devolver unos
# pocos caracteres de basura (números de página, un pie de imprenta) en vez de vacío, así que
# "hay algo" no sirve como criterio y hace falta un piso.
MINIMO_UTIL = 200


@dataclass(frozen=True)
class TextoCv:
    """`texto` y `warning` son INDEPENDIENTES y pueden coexistir: un CV largo tiene los dos."""
    texto: Optional[str] = None
    warning: Optional[str] = None


def extraer(contenido: bytes, filename: str) -> TextoCv:
    """Texto plano del CV, o el motivo por el que no se pudo. NUNCA levanta.

    Args:
        contenido: bytes del archivo, ya validados por `cv_service.validar`.
        filename: se usa para elegir el parser (la extensión ya viene validada).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            crudo, aviso = _de_pdf(contenido)
        elif ext == "docx":
            crudo, aviso = _de_docx(contenido), None
        elif ext == "doc":
            return TextoCv(warning="Formato .doc no soportado: pedile el CV en PDF o DOCX.")
        else:
            return TextoCv(warning=f"Formato .{ext} no soportado para extraer texto.")
    except Exception as exc:  # noqa: BLE001 — ver el encabezado: un archivo roto no corta el lote
        # 🔴 `archivo` y NO `filename`: `filename` es un atributo RESERVADO de `LogRecord` y
        # pasarlo en `extra` levanta KeyError — dentro del manejador de errores, o sea justo
        # cuando algo ya salió mal. Lo mismo con `module`, `lineno`, `funcName` y `process`.
        logger.warning("No se pudo extraer el texto de un CV",
                       extra={"archivo": filename, "error": str(exc)})
        return TextoCv(warning="El archivo está corrupto o no se pudo leer.")

    limpio = " ".join((crudo or "").split())
    if len(limpio) < MINIMO_UTIL:
        # 🔴 EL AVISO ESPECÍFICO GANA SOBRE EL GENÉRICO. Un PDF cifrado devuelve texto vacío, así
        # que cae en esta rama igual que un escaneo — y si el genérico lo pisara, el motivo se
        # perdería justo en el caso en que más sirve: "pedile la contraseña" y "pedile el CV en
        # otro formato" son acciones distintas. Es la razón por la que el warning es texto.
        return TextoCv(warning=aviso or "El archivo no tiene texto extraíble (probablemente sea "
                                        "un escaneo). Abrilo para leerlo.")
    if len(limpio) > MAX_CARACTERES:
        return TextoCv(texto=limpio[:MAX_CARACTERES],
                       warning=f"El CV es muy largo: se procesaron los primeros "
                               f"{MAX_CARACTERES:,} caracteres.".replace(",", "."))
    return TextoCv(texto=limpio, warning=aviso)


def _de_pdf(contenido: bytes) -> tuple:
    """(texto, aviso). El cifrado y las páginas ilegibles se reportan, no se propagan."""
    from pypdf import PdfReader

    lector = PdfReader(io.BytesIO(contenido))
    if lector.is_encrypted:
        # Muchos PDF "protegidos" solo tienen restricciones de impresión y abren con contraseña
        # vacía: se intenta, y recién si falla se marca. Descartarlos de entrada perdería CVs que
        # se podían leer perfectamente.
        try:
            if not lector.decrypt(""):
                raise ValueError("contraseña requerida")
        except Exception:
            return "", "El archivo está protegido con contraseña."
    if not lector.pages:
        return "", None
    partes, fallidas = [], 0
    for pagina in lector.pages:
        # Una página rota no invalida las otras: un CV de 3 páginas con la segunda corrupta sigue
        # siendo clasificable con las otras dos.
        try:
            partes.append(pagina.extract_text() or "")
        except Exception:  # noqa: BLE001
            fallidas += 1
    aviso = (f"No se pudieron leer {fallidas} página(s) del PDF." if fallidas else None)
    return "\n".join(partes), aviso


def _de_docx(contenido: bytes) -> str:
    """Párrafos + celdas de tabla: en un CV, media hoja de vida suele estar en una tabla."""
    from docx import Document

    doc = Document(io.BytesIO(contenido))
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)
